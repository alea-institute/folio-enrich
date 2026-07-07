from __future__ import annotations

import base64
import io
from typing import TYPE_CHECKING

from app.models.document import DocumentInput
from app.services.ingestion.base import IngestorBase

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


class WordIngestor(IngestorBase):
    # WordprocessingML main namespace
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def ingest(self, doc: DocumentInput) -> str:
        from docx import Document

        # Content is base64-encoded docx bytes
        try:
            docx_bytes = base64.b64decode(doc.content)
            document = Document(io.BytesIO(docx_bytes))
        except Exception:
            # Might be a file path
            document = Document(doc.content)

        return self._extract_all_text(document)

    def _extract_all_text(self, document: "DocxDocument") -> str:
        """Gather every ``w:t`` text node, grouped by paragraph, in document order.

        Iterating python-docx ``document.paragraphs`` skips paragraphs nested
        inside ``w:sdt``/``w:sdtContent`` (block-level content controls), so
        documents that wrap their body in content controls extract as empty
        (and the pipeline then mis-detects the raw ZIP bytes as plain text).
        Walking the parsed ``word/document.xml`` tree for all ``w:p`` elements —
        ``iter`` recurses through ``w:sdt``, tables, and text boxes — captures
        that text regardless of nesting.

        Each ``w:p`` is visited once by the top-level ``iter``; ``_para_text``
        prunes any *nested* ``w:p`` (e.g. a text box's ``w:txbxContent`` inside a
        run) so its text is emitted once — by its own visit — instead of being
        both glued onto the host paragraph and re-emitted (which corrupted output
        for docs containing text boxes / pull-quotes).

        NOTE: table-cell text (``w:tbl/w:tc/w:p``) is now included (it was
        excluded by ``document.paragraphs``). This is deliberate — legal
        documents carry substantive text in tables.
        """
        w = self._W_NS
        body = document.element.body  # lxml element for <w:body>

        paragraphs: list[str] = []
        for para in body.iter(f"{{{w}}}p"):
            text = self._para_text(para)
            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def _para_text(self, para) -> str:
        """Collect a paragraph's text, skipping any descendant ``w:p``.

        Descendant paragraphs (text boxes, nested content) are handled by their
        own top-level visit, so pruning them here avoids double emission and the
        boundary-merge that glued nested text onto the host paragraph.
        """
        w = self._W_NS
        parts: list[str] = []

        def walk(el) -> None:
            for child in el:
                tag = child.tag
                if tag == f"{{{w}}}p":
                    continue  # nested paragraph handled by its own visit
                if tag == f"{{{w}}}t":
                    if child.text:
                        parts.append(child.text)
                elif tag == f"{{{w}}}tab":
                    parts.append("\t")
                elif tag in (f"{{{w}}}br", f"{{{w}}}cr"):
                    parts.append("\n")
                else:
                    walk(child)

        walk(para)
        return "".join(parts).strip()
