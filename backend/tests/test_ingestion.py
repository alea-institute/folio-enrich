import base64
import io

import pytest

from app.models.document import DocumentFormat, DocumentInput
from app.services.ingestion.pdf_ingestor import PDFIngestor
from app.services.ingestion.plain_text import PlainTextIngestor
from app.services.ingestion.registry import detect_format, ingest


class TestPlainTextIngestor:
    def test_ingest_returns_content(self):
        doc = DocumentInput(content="Hello world")
        ingestor = PlainTextIngestor()
        assert ingestor.ingest(doc) == "Hello world"

    def test_ingest_preserves_whitespace(self):
        doc = DocumentInput(content="Line 1\n\nLine 2")
        ingestor = PlainTextIngestor()
        assert ingestor.ingest(doc) == "Line 1\n\nLine 2"


class TestFormatDetection:
    @pytest.mark.parametrize("filename,content,expected", [
        ("doc.txt", "", DocumentFormat.PLAIN_TEXT),
        ("readme.md", "", DocumentFormat.MARKDOWN),
        ("page.html", "", DocumentFormat.HTML),
        ("page.htm", "", DocumentFormat.HTML),
        ("file.pdf", "", DocumentFormat.PDF),
        ("file.docx", "", DocumentFormat.WORD),
        (None, "<html><body>Hi</body></html>", DocumentFormat.HTML),
        (None, "# Heading\n\nSome text", DocumentFormat.MARKDOWN),
        (None, "Just plain text.", DocumentFormat.PLAIN_TEXT),
    ])
    def test_format_detection(self, filename, content, expected):
        assert detect_format(filename, content) == expected


class TestIngestRegistry:
    def test_ingest_plain_text(self):
        doc = DocumentInput(content="test content", format=DocumentFormat.PLAIN_TEXT)
        assert ingest(doc) == "test content"

    def test_ingest_invalid_pdf_raises(self):
        doc = DocumentInput(content="test", format=DocumentFormat.PDF)
        with pytest.raises(Exception):
            ingest(doc)


# Minimal valid 1-page PDF with "Hello World" text
_MINI_PDF_B64 = (
    "JVBERi0xLjQKMSAwIG9iago8PCAvVHlwZSAvQ2F0YWxvZyAvUGFnZXMgMiAwIFIgPj4KZW5k"
    "b2JqCjIgMCBvYmoKPDwgL1R5cGUgL1BhZ2VzIC9LaWRzIFszIDAgUl0gL0NvdW50IDEgPj4K"
    "ZW5kb2JqCjMgMCBvYmoKPDwgL1R5cGUgL1BhZ2UgL1BhcmVudCAyIDAgUiAvTWVkaWFCb3gg"
    "WzAgMCA2MTIgNzkyXSAvQ29udGVudHMgNCAwIFIgL1Jlc291cmNlcyA8PCAvRm9udCA8PCAv"
    "RjEgNSAwIFIgPj4gPj4gPj4KZW5kb2JqCjQgMCBvYmoKPDwgL0xlbmd0aCA0NCA+PgpzdHJl"
    "YW0KQlQgL0YxIDEyIFRmIDEwMCA3MDAgVGQgKEhlbGxvIFdvcmxkKSBUaiBFVAplbmRzdHJl"
    "YW0KZW5kb2JqCjUgMCBvYmoKPDwgL1R5cGUgL0ZvbnQgL1N1YnR5cGUgL1R5cGUxIC9CYXNl"
    "Rm9udCAvSGVsdmV0aWNhID4+CmVuZG9iagp4cmVmCjAgNgowMDAwMDAwMDAwIDY1NTM1IGYg"
    "CjAwMDAwMDAwMDkgMDAwMDAgbiAKMDAwMDAwMDA1OCAwMDAwMCBuIAowMDAwMDAwMTE1MDAwMD"
    "AgbiAKMDAwMDAwMDI2NiAwMDAwMCBuIAowMDAwMDAwMzYwIDAwMDAwIG4gCnRyYWlsZXIKPDwg"
    "L1NpemUgNiAvUm9vdCAxIDAgUiA+PgpzdGFydHhyZWYKNDQxCiUlRU9G"
)


class TestPDFIngestor:
    """PDF text extraction via pypdf (pure-Python, BSD). PyMuPDF (AGPL) removed."""

    def test_ingest_extracts_text(self):
        doc = DocumentInput(content=_MINI_PDF_B64, filename="test.pdf")
        ingestor = PDFIngestor()
        text = ingestor.ingest(doc)
        assert "Hello World" in text

    def test_ingest_with_elements_returns_elements(self):
        doc = DocumentInput(content=_MINI_PDF_B64, filename="test.pdf")
        ingestor = PDFIngestor()
        text, elements = ingestor.ingest_with_elements(doc)
        assert "Hello World" in text
        assert len(elements) >= 1
        assert elements[0].page == 1

    def test_ingest_from_stream_without_pdf_extension(self):
        doc = DocumentInput(content=_MINI_PDF_B64, filename="test.bin")
        ingestor = PDFIngestor()
        text = ingestor.ingest(doc)
        assert "Hello World" in text


@pytest.mark.asyncio
class TestExtractEndpoint:
    """The /enrich/extract endpoint backs the upload UI: a base64 binary
    document in, readable plain text out (never the raw base64 blob)."""

    async def test_extract_pdf_returns_plain_text(self, client):
        r = await client.post(
            "/enrich/extract",
            json={"content": _MINI_PDF_B64, "filename": "test.pdf"},
        )
        assert r.status_code == 200
        body = r.json()
        assert "Hello World" in body["text"]
        assert body["format"] == "pdf"
        assert not body["text"].startswith("JVBER")  # not the base64 blob

    async def test_extract_passes_through_plain_text(self, client):
        r = await client.post(
            "/enrich/extract",
            json={"content": "Plain contract text.", "filename": "memo.txt"},
        )
        assert r.status_code == 200
        assert r.json()["text"] == "Plain contract text."

    async def test_extract_invalid_pdf_returns_422(self, client):
        r = await client.post(
            "/enrich/extract",
            json={"content": "not-a-real-base64-pdf", "filename": "broken.pdf"},
        )
        assert r.status_code == 422

    async def test_extract_html_strips_markup(self, client):
        html = (
            "<html><head><style>p{color:red}</style></head>"
            "<body><h1>Lease</h1><p>Tenant shall pay rent.</p></body></html>"
        )
        r = await client.post(
            "/enrich/extract", json={"content": html, "filename": "page.html"}
        )
        assert r.status_code == 200
        text = r.json()["text"]
        assert "Tenant shall pay rent." in text
        assert "<" not in text and "color:red" not in text  # tags + CSS gone

    async def test_extract_eml_parses_headers_and_body(self, client):
        eml = (
            "From: alice@example.com\nTo: bob@example.com\n"
            "Subject: Settlement Offer\nDate: Mon, 1 Jan 2024 10:00:00 -0000\n\n"
            "We propose to settle for $50,000.\n"
        )
        r = await client.post(
            "/enrich/extract", json={"content": eml, "filename": "mail.eml"}
        )
        assert r.status_code == 200
        text = r.json()["text"]
        assert "Subject: Settlement Offer" in text
        assert "We propose to settle for $50,000." in text

    async def test_extract_rtf_strips_control_codes(self, client):
        rtf = (
            r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times;}}"
            r"\f0\fs24 Arbitration shall occur in New York.\par}"
        )
        b64 = base64.b64encode(rtf.encode()).decode()
        r = await client.post(
            "/enrich/extract", json={"content": b64, "filename": "doc.rtf"}
        )
        assert r.status_code == 200
        text = r.json()["text"]
        assert "Arbitration shall occur in New York." in text
        assert "\\rtf" not in text and "fonttbl" not in text

    async def test_extract_docx_returns_paragraph_text(self, client):
        import io

        from docx import Document

        d = Document()
        d.add_paragraph("CONFIDENTIALITY AGREEMENT")
        d.add_paragraph("The Receiving Party shall protect Confidential Information.")
        buf = io.BytesIO()
        d.save(buf)
        b64 = base64.b64encode(buf.getvalue()).decode()
        r = await client.post(
            "/enrich/extract", json={"content": b64, "filename": "nda.docx"}
        )
        assert r.status_code == 200
        text = r.json()["text"]
        assert "CONFIDENTIALITY AGREEMENT" in text
        assert "Confidential Information" in text

    async def test_extract_markdown_strips_formatting(self, client):
        md = "# Master Agreement\n\nAcme **Corp** agrees to *indemnify* Beta LLC."
        r = await client.post(
            "/enrich/extract", json={"content": md, "filename": "note.md"}
        )
        assert r.status_code == 200
        text = r.json()["text"]
        assert "Master Agreement" in text
        assert "**" not in text and "#" not in text


class TestWordSdtContentControls:
    """Regression for B2: body text wrapped in <w:sdt> content controls.

    python-docx ``document.paragraphs`` does not descend into
    ``w:sdt``/``w:sdtContent``, so such documents used to extract 0 chars and
    the pipeline mis-detected the raw ZIP bytes as plain text. Damien's Trial
    Advocacy chapters wrap their body in content controls.
    """

    def _make_sdt_docx(self, text: str) -> str:
        import io

        from docx import Document
        from docx.oxml import OxmlElement

        d = Document()
        para = d.add_paragraph(text)
        body = d.element.body
        p_elem = para._p
        sdt = OxmlElement("w:sdt")
        sdt_content = OxmlElement("w:sdtContent")
        body.replace(p_elem, sdt)
        sdt.append(sdt_content)
        sdt_content.append(p_elem)  # paragraph now lives inside the content control
        buf = io.BytesIO()
        d.save(buf)
        return base64.b64encode(buf.getvalue()).decode()

    def test_word_ingestor_reads_sdt_wrapped_text(self):
        from app.models.document import DocumentInput
        from app.services.ingestion.word_ingestor import WordIngestor

        b64 = self._make_sdt_docx("Impeach the witness with the prior inconsistent statement.")
        text = WordIngestor().ingest(DocumentInput(content=b64))
        assert "Impeach the witness with the prior inconsistent statement." in text

    async def test_extract_route_reads_sdt_wrapped_text(self, client):
        b64 = self._make_sdt_docx("Object to leading questions on direct examination.")
        r = await client.post(
            "/enrich/extract", json={"content": b64, "filename": "chapter.docx"}
        )
        assert r.status_code == 200
        assert "Object to leading questions on direct examination." in r.json()["text"]


class TestWordNestedParagraphs:
    """Regression: text boxes / nested w:p must not double-count (B2 review)."""

    def _make_nested_p_docx(self, outer: str, nested: str) -> str:
        import io

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        d = Document()
        para = d.add_paragraph(outer)
        # Inject a nested <w:p> (as a text box would) inside the outer paragraph.
        nested_p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = nested
        r.append(t)
        nested_p.append(r)
        para._p.append(nested_p)
        buf = io.BytesIO()
        d.save(buf)
        return base64.b64encode(buf.getvalue()).decode()

    def test_nested_paragraph_text_appears_exactly_once(self):
        from app.models.document import DocumentInput
        from app.services.ingestion.word_ingestor import WordIngestor

        b64 = self._make_nested_p_docx("OUTER_TEXT", "TEXTBOX_TEXT")
        text = WordIngestor().ingest(DocumentInput(content=b64))
        assert text.count("TEXTBOX_TEXT") == 1
        assert text.count("OUTER_TEXT") == 1
        # And the two are not glued together into one token.
        assert "OUTER_TEXTTEXTBOX_TEXT" not in text
